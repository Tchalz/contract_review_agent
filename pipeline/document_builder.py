"""
Generates downloadable output documents from a completed contract review
report:

  - The "updated contract" — the original contract text with every
    human-approved clause rewording swapped in, clean (no markup).
  - The "change summary" — a standalone list of every approved edit:
    clause type, risk level, original wording, new wording, and the
    reason it was flagged.

Each is available as both .docx and .pdf. This module has no dependency
on the LLM/MCP layer — it works purely from the final report's flags and
the original contract_text already held by the caller (e.g. Streamlit's
session state), so producing these documents costs nothing beyond
formatting.
"""

import xml.sax.saxutils as saxutils
from datetime import date
from io import BytesIO

from docx import Document
from docx.shared import RGBColor

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

ACCENT_RGB = (216, 90, 48)  # matches the app's terracotta accent (#D85A30)
ACCENT_COLOR = colors.Color(216 / 255, 90 / 255, 48 / 255)
MUTED_COLOR = colors.Color(110 / 255, 110 / 255, 110 / 255)
LEVEL_ORDER = {"high": 0, "medium": 1, "low": 2}


def get_approved_changes(flags: list) -> list:
    """
    Returns only the flags that were human-approved AND have a suggested
    rewording that actually differs from the original snippet — i.e. the
    edits that should really be applied to the updated contract. Flags
    that were never surfaced for review (low risk) or were left
    unapproved are excluded.
    """
    changes = []
    for f in flags:
        if not f.get("human_approved"):
            continue
        rewording = (f.get("suggested_rewording") or "").strip()
        snippet = (f.get("snippet") or "").strip()
        if rewording and rewording != snippet:
            changes.append(f)
    return changes


def build_updated_contract_text(contract_text: str, flags: list) -> str:
    """
    Returns the contract text with every approved rewording swapped in
    for its original snippet. Unapproved / low-risk clauses are left
    exactly as written. Each snippet is a verbatim substring of
    contract_text (see mcp_server/app.py's _find_clause_snippet), so a
    plain single-occurrence replace is safe and exact.
    """
    updated = contract_text
    for change in get_approved_changes(flags):
        snippet = change["snippet"]
        rewording = change["suggested_rewording"]
        if snippet in updated:
            updated = updated.replace(snippet, rewording, 1)
    return updated


def _pdf_text(text: str) -> str:
    """reportlab's Paragraph treats its text as a small XML dialect, so
    any literal &, <, or > in the source contract must be escaped or the
    build breaks / silently drops content."""
    return saxutils.escape(text or "")


def _pdf_para_html(text: str) -> str:
    """Escapes text for a reportlab Paragraph AND converts embedded plain
    newlines into explicit <br/> tags. python-docx does this conversion
    automatically for us (a Word run with '\\n' in it renders as a real
    line break), but reportlab's Paragraph does not — without this, a
    clause heading like '1. LIABILITY\\nProvider shall...' runs together
    onto one line in the PDF instead of matching the docx output."""
    return _pdf_text(text).replace("\n", "<br/>")


_styles = getSampleStyleSheet()
_STYLE_TITLE = ParagraphStyle("CGTitle", parent=_styles["Title"], fontName="Helvetica-Bold", fontSize=20, spaceAfter=4)
_STYLE_META = ParagraphStyle("CGMeta", parent=_styles["Normal"], fontName="Helvetica-Oblique", fontSize=9, textColor=MUTED_COLOR, spaceAfter=2)
_STYLE_DISCLAIMER = ParagraphStyle("CGDisclaimer", parent=_styles["Normal"], fontName="Helvetica-Oblique", fontSize=9.5, textColor=MUTED_COLOR, spaceAfter=14, leading=13)
_STYLE_BODY = ParagraphStyle("CGBody", parent=_styles["Normal"], fontName="Helvetica", fontSize=11, leading=16, spaceAfter=10)
_STYLE_H2 = ParagraphStyle("CGH2", parent=_styles["Heading2"], fontName="Helvetica-Bold", fontSize=13, textColor=ACCENT_COLOR, spaceBefore=14, spaceAfter=6)
_STYLE_LABEL = ParagraphStyle("CGLabel", parent=_styles["Normal"], fontName="Helvetica", fontSize=10.5, leading=14, spaceAfter=8)


def _disclaimer(num_changes: int) -> str:
    return (
        f"{num_changes} clause(s) updated based on human-approved suggestions during review. "
        "Any clauses not listed in the accompanying change summary are unchanged from the "
        "original document. This is not legal advice — review the full document before signing."
    )


# ---------------------------------------------------------------------------
# Updated contract
# ---------------------------------------------------------------------------

def build_updated_contract_docx(contract_text: str, flags: list, title: str) -> bytes:
    updated_text = build_updated_contract_text(contract_text, flags)
    changes = get_approved_changes(flags)

    doc = Document()
    doc.add_heading(title, level=0)

    p = doc.add_paragraph()
    run = p.add_run(f"Updated {date.today().strftime('%B %d, %Y')} — generated by Clausegraph")
    run.italic = True

    p = doc.add_paragraph()
    run = p.add_run(_disclaimer(len(changes)))
    run.italic = True

    doc.add_paragraph()  # spacer

    for para in updated_text.split("\n\n"):
        para = para.strip()
        if para:
            doc.add_paragraph(para)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_updated_contract_pdf(contract_text: str, flags: list, title: str) -> bytes:
    updated_text = build_updated_contract_text(contract_text, flags)
    changes = get_approved_changes(flags)

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=LETTER,
        topMargin=0.85 * inch, bottomMargin=0.85 * inch,
        leftMargin=0.9 * inch, rightMargin=0.9 * inch,
    )
    story = [
        Paragraph(_pdf_text(title), _STYLE_TITLE),
        Paragraph(_pdf_text(f"Updated {date.today().strftime('%B %d, %Y')} — generated by Clausegraph"), _STYLE_META),
        Paragraph(_pdf_text(_disclaimer(len(changes))), _STYLE_DISCLAIMER),
    ]
    for para in updated_text.split("\n\n"):
        para = para.strip()
        if para:
            story.append(Paragraph(_pdf_para_html(para), _STYLE_BODY))

    doc.build(story)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Change summary
# ---------------------------------------------------------------------------

def build_change_summary_docx(flags: list, title: str) -> bytes:
    changes = get_approved_changes(flags)
    doc = Document()

    doc.add_heading(f"{title} — Change Summary", level=0)
    p = doc.add_paragraph()
    run = p.add_run(f"Generated {date.today().strftime('%B %d, %Y')} by Clausegraph")
    run.italic = True
    doc.add_paragraph()

    if not changes:
        doc.add_paragraph("No clauses were approved for rewording during human review.")
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    for change in sorted(changes, key=lambda f: LEVEL_ORDER.get(f["level"], 3)):
        doc.add_heading(
            f"{change['level'].upper()} — {change['clause_type'].replace('_', ' ').title()}", level=2
        )

        p = doc.add_paragraph()
        p.add_run("Reason flagged: ").bold = True
        p.add_run(change.get("reason", ""))

        p = doc.add_paragraph()
        p.add_run("Original: ").bold = True
        p.add_run(change.get("snippet", ""))

        p = doc.add_paragraph()
        p.add_run("Approved new wording: ").bold = True
        new_run = p.add_run(change.get("suggested_rewording", ""))
        new_run.font.color.rgb = RGBColor(*ACCENT_RGB)

        doc.add_paragraph()  # spacer

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_change_summary_pdf(flags: list, title: str) -> bytes:
    changes = get_approved_changes(flags)

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=LETTER,
        topMargin=0.85 * inch, bottomMargin=0.85 * inch,
        leftMargin=0.9 * inch, rightMargin=0.9 * inch,
    )
    story = [
        Paragraph(_pdf_text(f"{title} — Change Summary"), _STYLE_TITLE),
        Paragraph(_pdf_text(f"Generated {date.today().strftime('%B %d, %Y')} by Clausegraph"), _STYLE_META),
        Spacer(1, 10),
    ]

    if not changes:
        story.append(Paragraph("No clauses were approved for rewording during human review.", _STYLE_BODY))
        doc.build(story)
        return buf.getvalue()

    for change in sorted(changes, key=lambda f: LEVEL_ORDER.get(f["level"], 3)):
        story.append(Paragraph(
            _pdf_text(f"{change['level'].upper()} — {change['clause_type'].replace('_', ' ').title()}"),
            _STYLE_H2,
        ))
        story.append(Paragraph(f"<b>Reason flagged:</b> {_pdf_text(change.get('reason', ''))}", _STYLE_LABEL))
        story.append(Paragraph(f"<b>Original:</b> {_pdf_text(change.get('snippet', ''))}", _STYLE_LABEL))
        story.append(Paragraph(
            f"<b>Approved new wording:</b> "
            f"<font color='#{ACCENT_RGB[0]:02x}{ACCENT_RGB[1]:02x}{ACCENT_RGB[2]:02x}'>"
            f"{_pdf_text(change.get('suggested_rewording', ''))}</font>",
            _STYLE_LABEL,
        ))

    doc.build(story)
    return buf.getvalue()
